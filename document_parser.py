"""
document_parser.py —— 文档解析器（txt / md / docx / pdf）
==========================================================

支持格式：
    .txt  - UTF-8 纯文本
    .md   - Markdown（含 YAML frontmatter 解析）
    .docx - Microsoft Word（用 python-docx 提段落）
    .pdf  - PDF（用 pypdf 提所有页文本）

所有解析器返回统一的 tuple：
    (text: str, metadata: dict)

metadata 通常包含：
    - title: str        从 frontmatter 提取（可选）
    - tags: list        从 frontmatter 提取（可选）
    - source: str       从 frontmatter 提取（可选）
    - page_count: int   PDF / DOCX 的页/段落数（可选）
    - char_count: int   文本字符数（计算）

限制：
    - 单文件最大 20MB（在 endpoint 校验）
    - DOCX 只提段落文本，不提表格、批注、图片
    - PDF 只提文本层，不做 OCR（扫描版 PDF 提不到内容）
    - Markdown frontmatter 必须是 YAML（用简单正则解析，不引 pyyaml 依赖）
"""

import re
from typing import Tuple, Dict, List


# ============== Markdown frontmatter ==============

# 匹配文件开头的 YAML frontmatter：
#   ---
#   title: "..."
#   tags: [a, b]
#   ---
FRONTMATTER_RE = re.compile(
    r'^---\s*\n(.*?)\n---\s*\n(.*)$',
    re.DOTALL,
)


def _parse_simple_yaml(text: str) -> Dict:
    """
    极简 YAML 解析（只支持 key: value 形式 + key: [list]）。
    不依赖 pyyaml，避免额外装包。
    """
    result: Dict = {}
    for line in text.split('\n'):
        line = line.rstrip()
        if not line or ':' not in line:
            continue
        key, _, value = line.partition(':')
        key = key.strip()
        value = value.strip()
        # 去掉引号
        if (value.startswith('"') and value.endswith('"')) or \
           (value.startswith("'") and value.endswith("'")):
            value = value[1:-1]
        # 列表 [a, b, c]
        if value.startswith('[') and value.endswith(']'):
            inner = value[1:-1].strip()
            items = [i.strip().strip('"\'') for i in inner.split(',') if i.strip()]
            result[key] = items
        # 布尔 / 数字
        elif value.lower() in ('true', 'false'):
            result[key] = value.lower() == 'true'
        else:
            try:
                result[key] = int(value)
            except ValueError:
                try:
                    result[key] = float(value)
                except ValueError:
                    result[key] = value
    return result


def parse_markdown(content: bytes) -> Tuple[str, Dict]:
    """
    解析 Markdown，提取 frontmatter。

    Returns:
        (text, metadata) —— text 是去掉 frontmatter 后的 markdown 正文
    """
    try:
        text = content.decode('utf-8')
    except UnicodeDecodeError:
        text = content.decode('utf-8', errors='replace')

    m = FRONTMATTER_RE.match(text)
    metadata: Dict = {}
    if m:
        yaml_text = m.group(1)
        text = m.group(2).lstrip('\n')
        metadata = _parse_simple_yaml(yaml_text)

    return text, metadata


# ============== TXT ==============

def parse_txt(content: bytes) -> Tuple[str, Dict]:
    """
    解析纯文本。尝试 UTF-8，失败用 GBK（兼容 Windows .txt 默认编码）。
    """
    for encoding in ('utf-8', 'utf-8-sig', 'gbk', 'gb18030'):
        try:
            return content.decode(encoding), {}
        except UnicodeDecodeError:
            continue
    # 兜底：替换非法字符
    return content.decode('utf-8', errors='replace'), {}


# ============== DOCX ==============

def parse_docx(content: bytes) -> Tuple[str, Dict]:
    """
    解析 .docx，用 python-docx 提取段落文本。
    表格 / 图片 / 批注 不提取（够用起步）。
    """
    from io import BytesIO
    from docx import Document

    doc = Document(BytesIO(content))
    paragraphs: List[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # 提取表格内文本（可选）
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text and cell_text not in paragraphs:
                    paragraphs.append(cell_text)

    text = '\n\n'.join(paragraphs)
    metadata = {
        "para_count": len(paragraphs),
        "has_tables": len(doc.tables) > 0,
    }

    # core_properties 可能含标题（Word 文件元数据）
    if doc.core_properties.title:
        metadata["doc_title"] = doc.core_properties.title

    return text, metadata


# ============== PDF ==============

def parse_pdf(content: bytes) -> Tuple[str, Dict]:
    """
    解析 .pdf，用 pypdf 提取所有页文本。
    不做 OCR（扫描版 PDF 提不到内容，会返回空）。
    """
    from io import BytesIO
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(content))
    page_texts: List[str] = []
    for i, page in enumerate(reader.pages):
        try:
            page_text = page.extract_text() or ""
        except Exception as e:
            page_text = f"[第 {i+1} 页解析失败: {e}]"
        if page_text.strip():
            page_texts.append(page_text)

    # 用 "\n\n" 隔开各页（保留页边界）
    text = "\n\n".join(page_texts)
    metadata = {
        "page_count": len(reader.pages),
        "extracted_pages": len(page_texts),
    }

    # PDF 元数据
    if reader.metadata:
        if reader.metadata.title:
            metadata["pdf_title"] = reader.metadata.title
        if reader.metadata.author:
            metadata["author"] = reader.metadata.author

    return text, metadata


# ============== 统一分发 ==============

# 文件扩展名 → 解析器
PARSERS = {
    ".txt": parse_txt,
    ".md": parse_markdown,
    ".markdown": parse_markdown,
    ".docx": parse_docx,
    ".pdf": parse_pdf,
}


def parse_document(filename: str, content: bytes) -> Tuple[str, Dict]:
    """
    统一入口：根据文件名后缀分发到对应解析器。

    Args:
        filename: 文件名（用于判断扩展名）
        content: 文件二进制内容

    Returns:
        (text, metadata)

    Raises:
        ValueError: 不支持的文件类型
        Exception: 解析失败（具体异常透传）
    """
    # 取扩展名（小写）
    ext = "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    parser = PARSERS.get(ext)
    if not parser:
        supported = ", ".join(sorted(PARSERS.keys()))
        raise ValueError(f"不支持的文件类型: {ext}（支持: {supported}）")

    text, metadata = parser(content)
    # 统一附加元数据
    metadata["filename"] = filename
    metadata["extension"] = ext
    metadata["char_count"] = len(text)
    metadata["byte_count"] = len(content)

    return text, metadata


def is_supported_file(filename: str) -> bool:
    """快速判断是否支持的文件类型"""
    if "." not in filename:
        return False
    ext = "." + filename.rsplit(".", 1)[-1].lower()
    return ext in PARSERS


def supported_extensions() -> List[str]:
    """返回支持的文件扩展名列表"""
    return sorted(PARSERS.keys())